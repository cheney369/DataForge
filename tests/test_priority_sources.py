from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from dataforge.application import DataForge
from dataforge.config import Settings
from dataforge.ingestion import preview_source_records


def write_text_pdf(path: Path, text: str) -> None:
    writer = PdfWriter()
    page = writer.add_blank_page(width=420, height=595)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
    )
    content = DecodedStreamObject()
    content.set_data(f"BT /F1 12 Tf 50 520 Td ({text}) Tj ET".encode("ascii"))
    page[NameObject("/Contents")] = writer._add_object(content)
    with path.open("wb") as handle:
        writer.write(handle)


class PrioritySourceFormatsTest(unittest.TestCase):
    def setUp(self):
        self.temporary = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary.name)
        self.app = DataForge(Settings.load(self.root))

    def tearDown(self):
        self.temporary.cleanup()

    def test_priority_formats_reach_published_assets(self):
        files: list[tuple[Path, str]] = []

        text_file = self.root / "follow-up.txt"
        text_file.write_text("患者血压稳定，继续随访。", encoding="utf-8")
        files.append((text_file, "患者血压稳定"))

        markdown_file = self.root / "guide.md"
        markdown_file.write_text("# 用药指南\n\n每日按时服药。", encoding="utf-8")
        files.append((markdown_file, "每日按时服药"))

        csv_file = self.root / "faq.csv"
        csv_file.write_text("question,answer\n如何预约,通过医院小程序预约\n", encoding="utf-8")
        files.append((csv_file, "通过医院小程序预约"))

        excel_file = self.root / "follow-up.xlsx"
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "随访计划"
        sheet.append(["患者", "计划"])
        sheet.append(["张三", "每周测量血压"])
        workbook.save(excel_file)
        files.append((excel_file, "每周测量血压"))

        word_file = self.root / "report.docx"
        document = Document()
        document.add_heading("出院记录", level=1)
        document.add_paragraph("患者恢复良好，可以出院。")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "复诊"
        table.cell(0, 1).text = "两周后"
        document.save(word_file)
        files.append((word_file, "患者恢复良好"))

        pdf_file = self.root / "english-report.pdf"
        write_text_pdf(pdf_file, "Medical follow-up is stable")
        files.append((pdf_file, "Medical follow-up is stable"))

        for source_file, expected_text in files:
            with self.subTest(suffix=source_file.suffix):
                result = self.app.flow(
                    source_file,
                    name=source_file.stem,
                    engine_override="native",
                )
                self.assertEqual(result.run["status"], "completed")
                self.assertEqual(result.asset_version["status"], "published")
                blob = self.app.blobs.resolve(result.asset_version["blob_uri"])
                records = [json.loads(line) for line in blob.read_text(encoding="utf-8").splitlines()]
                self.assertTrue(any(expected_text in item["content"] for item in records))

    def test_pdf_and_docx_preview_include_precise_source_positions(self):
        pdf_file = self.root / "located.pdf"
        write_text_pdf(pdf_file, "Follow-up page one")
        pdf_ingestion = self.app.sources.ingest(pdf_file)
        pdf_preview = preview_source_records(
            self.app.blobs.resolve(pdf_ingestion.source_version["blob_uri"]),
            pdf_ingestion.source_version,
        )
        self.assertEqual(pdf_preview["records"][0]["source_locator"]["page_number"], 1)

        word_file = self.root / "located.docx"
        document = Document()
        document.add_paragraph("第一段")
        document.add_paragraph("第二段")
        document.save(word_file)
        word_ingestion = self.app.sources.ingest(word_file)
        word_preview = preview_source_records(
            self.app.blobs.resolve(word_ingestion.source_version["blob_uri"]),
            word_ingestion.source_version,
        )
        self.assertEqual(
            [item["source_locator"]["paragraph_index"] for item in word_preview["records"]],
            [1, 2],
        )

    def test_xlsx_preview_preserves_sheet_and_row_position(self):
        excel_file = self.root / "located.xlsx"
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "门诊随访"
        sheet.append(["姓名", "复诊日期"])
        sheet.append(["李四", "2026-09-01"])
        workbook.save(excel_file)

        ingestion = self.app.sources.ingest(excel_file)
        preview = preview_source_records(
            self.app.blobs.resolve(ingestion.source_version["blob_uri"]),
            ingestion.source_version,
        )

        locator = preview["records"][0]["source_locator"]
        self.assertEqual(locator["kind"], "xlsx")
        self.assertEqual(locator["sheet_name"], "门诊随访")
        self.assertEqual(locator["sheet_index"], 1)
        self.assertEqual(locator["row_number"], 2)
        self.assertEqual(locator["header_row_number"], 1)


if __name__ == "__main__":
    unittest.main()
