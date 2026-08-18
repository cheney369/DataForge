from __future__ import annotations

import csv
import json
import mimetypes
from pathlib import Path
from typing import Any, Iterable

from .blobs import BlobStore
from .database import MetadataStore
from .errors import ValidationError
from .models import IngestResult


SUPPORTED_SUFFIXES = {".pdf", ".csv", ".xlsx", ".md", ".docx", ".txt", ".json", ".jsonl"}
CONTENT_FIELDS = ("raw_content", "content", "text", "body", "answer")


class SourceService:
    def __init__(self, store: MetadataStore, blobs: BlobStore):
        self.store = store
        self.blobs = blobs

    def ingest(
        self,
        file_path: str | Path,
        *,
        source_id: str | None = None,
        name: str | None = None,
        kind: str = "file",
        metadata: dict[str, Any] | None = None,
    ) -> IngestResult:
        path = Path(file_path).expanduser().resolve()
        if not path.is_file():
            raise ValidationError(f"Source file does not exist: {path}")
        if path.suffix.lower() not in SUPPORTED_SUFFIXES:
            supported = ", ".join(sorted(SUPPORTED_SUFFIXES))
            raise ValidationError(f"Unsupported source type {path.suffix!r}; supported: {supported}")

        blob_uri, sha256, size_bytes = self.blobs.put_file(path)
        source = (
            self.store.get_source(source_id)
            if source_id
            else self.store.create_source(name or path.stem, kind, metadata or {})
        )
        existing = self.store.find_source_version_by_hash(source["id"], sha256)
        if existing:
            return IngestResult(source=source, source_version=existing, created=False)

        media_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
        version = self.store.create_source_version(
            source_id=source["id"],
            blob_uri=blob_uri,
            sha256=sha256,
            size_bytes=size_bytes,
            media_type=media_type,
            original_filename=path.name,
        )
        return IngestResult(source=source, source_version=version, created=True)


def materialize_source_records(
    source_file: Path,
    source_version: dict[str, Any],
    destination: Path,
) -> int:
    suffix = Path(source_version["original_filename"]).suffix.lower()
    destination.parent.mkdir(parents=True, exist_ok=True)
    count = 0
    with destination.open("w", encoding="utf-8") as handle:
        for index, (record, locator) in enumerate(_read_records(source_file, suffix)):
            content = _extract_content(record)
            if not content.strip():
                continue
            materialized = {
                "document_id": f"{source_version['id']}:{index}",
                "source_id": source_version["source_id"],
                "source_version_id": source_version["id"],
                "source_record_index": index,
                "source_locator": locator,
                "raw_content": content,
            }
            handle.write(json.dumps(materialized, ensure_ascii=False) + "\n")
            count += 1

    if count == 0:
        raise ValidationError("Source records do not contain any text")
    return count


def preview_source_records(
    source_file: Path,
    source_version: dict[str, Any],
    *,
    max_records: int = 20,
    max_chars: int = 12_000,
) -> dict[str, Any]:
    """Return a bounded, text-first preview without materializing another asset."""
    suffix = Path(source_version["original_filename"]).suffix.lower()
    records: list[dict[str, Any]] = []
    character_count = 0
    truncated = False

    for index, (record, locator) in enumerate(_read_records(source_file, suffix)):
        content = _extract_content(record)
        if not content.strip():
            continue
        if len(records) >= max_records or character_count >= max_chars:
            truncated = True
            break

        remaining = max_chars - character_count
        excerpt = content[:remaining]
        records.append({"index": index, "content": excerpt, "source_locator": locator})
        character_count += len(excerpt)
        if len(excerpt) < len(content):
            truncated = True
            break

    if not records:
        raise ValidationError("Source records do not contain any text")

    return {
        "records": records,
        "preview_record_count": len(records),
        "character_count": character_count,
        "truncated": truncated,
    }


def _read_records(source_file: Path, suffix: str) -> Iterable[tuple[Any, dict[str, Any]]]:
    if suffix in {".txt", ".md"}:
        text = source_file.read_text(encoding="utf-8-sig", errors="replace")
        yield text, {
            "kind": "document",
            "character_start": 0,
            "character_end": len(text),
        }
        return
    if suffix == ".pdf":
        yield from _read_pdf(source_file)
        return
    if suffix == ".docx":
        yield from _read_docx(source_file)
        return
    if suffix == ".json":
        payload = json.loads(source_file.read_text(encoding="utf-8-sig"))
        if isinstance(payload, list):
            for index, item in enumerate(payload):
                yield item, {"kind": "json", "json_index": index}
        else:
            yield payload, {"kind": "json", "json_path": "$"}
        return
    if suffix == ".jsonl":
        with source_file.open(encoding="utf-8-sig") as handle:
            for line_number, line in enumerate(handle, start=1):
                if not line.strip():
                    continue
                try:
                    yield json.loads(line), {"kind": "jsonl", "line_number": line_number}
                except json.JSONDecodeError as exc:
                    raise ValidationError(f"Invalid JSONL at line {line_number}: {exc}") from exc
        return
    if suffix == ".csv":
        with source_file.open(encoding="utf-8-sig", newline="") as handle:
            for line_number, row in enumerate(csv.DictReader(handle), start=2):
                yield row, {"kind": "csv", "line_number": line_number}
        return
    if suffix == ".xlsx":
        yield from _read_xlsx(source_file)
        return
    raise ValidationError(f"No reader for source suffix: {suffix}")


def _read_pdf(source_file: Path) -> Iterable[tuple[str, dict[str, Any]]]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(source_file)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise ValidationError("PDF 文件已加密，无法读取") from exc
        for page_number, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception as exc:
                raise ValidationError(f"无法读取 PDF 第 {page_number} 页：{exc}") from exc
            if text.strip():
                yield text, {
                    "kind": "pdf",
                    "page_number": page_number,
                    "character_start": 0,
                    "character_end": len(text),
                }
    except ValidationError:
        raise
    except Exception as exc:
        raise ValidationError(f"无法解析 PDF 文件：{exc}") from exc


def _read_docx(source_file: Path) -> Iterable[tuple[str, dict[str, Any]]]:
    try:
        from docx import Document

        document = Document(source_file)
    except Exception as exc:
        raise ValidationError(f"无法解析 Word 文件：{exc}") from exc

    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        if paragraph.text.strip():
            text = paragraph.text.strip()
            yield text, {
                "kind": "docx_paragraph",
                "paragraph_index": paragraph_index,
                "character_start": 0,
                "character_end": len(text),
            }
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                text = "\t".join(values)
                yield text, {
                    "kind": "docx_table_row",
                    "table_index": table_index,
                    "row_index": row_index,
                    "character_start": 0,
                    "character_end": len(text),
                }


def _read_xlsx(source_file: Path) -> Iterable[tuple[dict[str, Any], dict[str, Any]]]:
    source_handle = None
    workbook = None
    try:
        from openpyxl import load_workbook

        source_handle = source_file.open("rb")
        workbook = load_workbook(source_handle, read_only=True, data_only=True)
    except Exception as exc:
        if source_handle is not None:
            source_handle.close()
        raise ValidationError(f"无法解析 Excel 文件：{exc}") from exc

    try:
        for sheet_index, worksheet in enumerate(workbook.worksheets, start=1):
            rows = worksheet.iter_rows(values_only=True)
            header_row_number = 0
            headers: list[str] = []
            for row_number, values in enumerate(rows, start=1):
                normalized = [_xlsx_value(value) for value in values]
                if not any(value not in {None, ""} for value in normalized):
                    continue
                header_row_number = row_number
                headers = _xlsx_headers(normalized)
                break

            if not headers:
                continue

            for row_number, values in enumerate(rows, start=header_row_number + 1):
                normalized = [_xlsx_value(value) for value in values]
                if not any(value not in {None, ""} for value in normalized):
                    continue
                record = {
                    header: normalized[index] if index < len(normalized) else None
                    for index, header in enumerate(headers)
                }
                yield record, {
                    "kind": "xlsx",
                    "sheet_index": sheet_index,
                    "sheet_name": worksheet.title,
                    "row_number": row_number,
                    "header_row_number": header_row_number,
                }
    finally:
        if workbook is not None:
            workbook.close()
        if source_handle is not None:
            source_handle.close()


def _xlsx_headers(values: list[Any]) -> list[str]:
    headers: list[str] = []
    used: dict[str, int] = {}
    for index, value in enumerate(values, start=1):
        base = str(value).strip() if value not in {None, ""} else f"column_{index}"
        occurrence = used.get(base, 0) + 1
        used[base] = occurrence
        headers.append(base if occurrence == 1 else f"{base}_{occurrence}")
    return headers


def _xlsx_value(value: Any) -> Any:
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    isoformat = getattr(value, "isoformat", None)
    if callable(isoformat):
        return isoformat()
    return str(value)


def _extract_content(record: Any) -> str:
    if isinstance(record, str):
        return record
    if isinstance(record, dict):
        for field in CONTENT_FIELDS:
            value = record.get(field)
            if isinstance(value, str):
                return value
        return json.dumps(record, ensure_ascii=False, sort_keys=True)
    return str(record)
